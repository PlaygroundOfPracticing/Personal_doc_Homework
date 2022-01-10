import re
import jieba
import docx
import os
import re
import time
import train
from win32com import client as wc
import time
startTime = time.time()

strSample = '今天天气不错，但是我很伤心。'
sample = train.predict_matrix(strSample, train.word_list[:-382])

# 记得修改路径
docPath = './语料/'
docxPath = './docx/'
docxPath2 = './docx/'
resultPath = './result'
pattern1 = re.compile(r'，')
pattern2 = re.compile(r'。')
pattern3 = re.compile(r'？')
pattern4 = re.compile(r'！')
pattern5 = re.compile(r'；')

# 将doc文件转为docx

# def doc2docx(docPath, docxPath):
#     for root, dirs, files in os.walk(docPath):
#         word = wc.Dispatch('Word.Application')
#         for file in files:
#             doc = word.Documents.Open(root + file)
#             doc.SaveAs(docxPath + file + 'x', 12)
#             doc.Close()
#         word.Quit()
#
#
# doc2docx(docPath, docxPath)

for i in range(100):
    if i < 9:
        doc = docx.Document(docxPath2 + '00' + str(i + 1) + '.docx')
        resultFile = str(resultPath + '00' + str(i + 1) + '.txt')
    elif i < 99:
        doc = docx.Document(docxPath2 + '0' + str(i + 1) + '.docx')
        resultFile = str(resultPath + '0' + str(i + 1) + '.txt')
    elif i == 99:
        doc = docx.Document(docxPath2 + '' + str(i + 1) + '.docx')
        resultFile = str(resultPath + '' + str(i + 1) + '.txt')

    for paragraph in doc.paragraphs[1:]:
        text_division = []
        if paragraph.text != '':
            sentence = paragraph.text
            sentence = re.sub(pattern1, '，|', sentence)
            sentence = re.sub(pattern2, '。|', sentence)
            sentence = re.sub(pattern3, '？|', sentence)
            sentence = re.sub(pattern4, '！|', sentence)
            sentence = re.sub(pattern5, '；|', sentence)
            text_list = sentence.split('|')[:-1]
            markNumber = 1
            for j in range(len(text_list)):
                if j == len(text_list)-1:
                    text_division.append(''.join([text_list[j], '']))
                elif text_list[j][-1] != '，':
                    markNumber += 1
                    text_division.append(''.join([text_list[j], str(markNumber)+'、']))
                elif train.clf.predict([train.predict_matrix(text_list[j], train.word_list[:-382])]) == 1:
                    markNumber += 1
                    text_division.append(''.join([text_list[j], str(markNumber)+'、']))
            with open(resultFile, 'w', encoding='utf-8') as f:
                f.write('1、'+''.join(text_division))
            # print('1、'+''.join(text_division))
print('运行时间：'+str(time.time()-startTime))

